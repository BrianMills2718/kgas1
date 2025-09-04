#!/usr/bin/env python3
"""
Document Ingestion System for Multiple Formats

Provides comprehensive document ingestion capabilities for PDF, Word, PowerPoint, 
Excel, HTML, Markdown, plain text, and other formats with intelligent content 
extraction, metadata preservation, structure detection, and quality assessment.
"""

import logging
import asyncio
import mimetypes
import hashlib
import json
import time
from typing import Dict, List, Optional, Any, Union, Tuple, Iterator, Callable
from dataclasses import dataclass, field, asdict
from datetime import datetime
from pathlib import Path
from contextlib import contextmanager
from abc import ABC, abstractmethod
from enum import Enum
import tempfile
import shutil
import zipfile
import tarfile
from urllib.parse import urlparse
import re

# Import extraction libraries
try:
    import PyPDF2
    import pdfplumber
    from pdfminer.high_level import extract_text as pdfminer_extract
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from pptx import Presentation
    PPTX_AVAILABLE = True
except ImportError:
    PPTX_AVAILABLE = False

try:
    import pandas as pd
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    import requests
    HTML_AVAILABLE = True
except ImportError:
    HTML_AVAILABLE = False

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False

try:
    import eml_parser
    EMAIL_AVAILABLE = True
except ImportError:
    EMAIL_AVAILABLE = False

logger = logging.getLogger(__name__)


class DocumentFormat(Enum):
    """Supported document formats"""
    PDF = "pdf"
    DOCX = "docx"
    DOC = "doc"
    PPTX = "pptx"
    PPT = "ppt"
    XLSX = "xlsx"
    XLS = "xls"
    CSV = "csv"
    HTML = "html"
    XML = "xml"
    MARKDOWN = "markdown"
    TXT = "txt"
    RTF = "rtf"
    JSON = "json"
    EMAIL = "email"
    EPUB = "epub"
    ODT = "odt"
    UNKNOWN = "unknown"


class ExtractionQuality(Enum):
    """Quality levels for content extraction"""
    EXCELLENT = "excellent"
    GOOD = "good"
    FAIR = "fair"
    POOR = "poor"
    FAILED = "failed"


@dataclass
class DocumentMetadata:
    """Comprehensive document metadata"""
    # Basic file information
    filename: str
    file_path: Optional[str] = None
    file_size_bytes: int = 0
    file_hash: Optional[str] = None
    mime_type: Optional[str] = None
    detected_format: DocumentFormat = DocumentFormat.UNKNOWN
    
    # Content information
    page_count: int = 0
    word_count: int = 0
    character_count: int = 0
    paragraph_count: int = 0
    
    # Document properties
    title: Optional[str] = None
    author: Optional[str] = None
    subject: Optional[str] = None
    creator: Optional[str] = None
    producer: Optional[str] = None
    keywords: List[str] = field(default_factory=list)
    
    # Timestamps
    created_date: Optional[datetime] = None
    modified_date: Optional[datetime] = None
    accessed_date: Optional[datetime] = None
    ingestion_date: datetime = field(default_factory=datetime.now)
    
    # Structure information
    has_tables: bool = False
    has_images: bool = False
    has_links: bool = False
    has_headers: bool = False
    has_footnotes: bool = False
    
    # Quality metrics
    extraction_quality: ExtractionQuality = ExtractionQuality.FAILED
    extraction_confidence: float = 0.0
    text_extraction_method: Optional[str] = None
    
    # Language and encoding
    detected_language: Optional[str] = None
    encoding: Optional[str] = None
    
    # Custom metadata
    custom_fields: Dict[str, Any] = field(default_factory=dict)


@dataclass
class ExtractedContent:
    """Extracted content from document"""
    # Text content
    full_text: str = ""
    structured_text: Dict[str, Any] = field(default_factory=dict)
    
    # Document structure
    paragraphs: List[str] = field(default_factory=list)
    headings: List[Dict[str, Any]] = field(default_factory=list)
    tables: List[Dict[str, Any]] = field(default_factory=list)
    images: List[Dict[str, Any]] = field(default_factory=list)
    links: List[Dict[str, Any]] = field(default_factory=list)
    footnotes: List[Dict[str, Any]] = field(default_factory=list)
    
    # Metadata
    metadata: DocumentMetadata = field(default_factory=DocumentMetadata)
    
    # Processing information
    extraction_time: float = 0.0
    processing_notes: List[str] = field(default_factory=list)
    
    def get_text_length(self) -> int:
        """Get total text length"""
        return len(self.full_text)
    
    def get_structure_summary(self) -> Dict[str, int]:
        """Get summary of document structure"""
        return {
            "paragraphs": len(self.paragraphs),
            "headings": len(self.headings),
            "tables": len(self.tables),
            "images": len(self.images),
            "links": len(self.links),
            "footnotes": len(self.footnotes)
        }


class DocumentExtractor(ABC):
    """Abstract base class for document extractors"""
    
    @abstractmethod
    def can_extract(self, file_path: Path, mime_type: str) -> bool:
        """Check if this extractor can handle the file"""
        pass
    
    @abstractmethod
    def extract_content(self, file_path: Path) -> ExtractedContent:
        """Extract content from document"""
        pass
    
    @abstractmethod
    def get_supported_formats(self) -> List[DocumentFormat]:
        """Get list of supported formats"""
        pass
    
    def get_extractor_info(self) -> Dict[str, Any]:
        """Get information about this extractor"""
        return {
            "extractor_name": self.__class__.__name__,
            "supported_formats": [f.value for f in self.get_supported_formats()],
            "dependencies_available": True
        }


class PDFExtractor(DocumentExtractor):
    """PDF document extractor with multiple extraction strategies"""
    
    def __init__(self):
        self.available_methods = []
        if PDF_AVAILABLE:
            self.available_methods = ["pdfplumber", "pypdf2", "pdfminer"]
    
    def can_extract(self, file_path: Path, mime_type: str) -> bool:
        """Check if this is a PDF file"""
        return (PDF_AVAILABLE and 
                (file_path.suffix.lower() == '.pdf' or 
                 mime_type in ['application/pdf']))
    
    def get_supported_formats(self) -> List[DocumentFormat]:
        """Get supported formats"""
        return [DocumentFormat.PDF] if PDF_AVAILABLE else []
    
    def extract_content(self, file_path: Path) -> ExtractedContent:
        """Extract content from PDF using best available method"""
        if not PDF_AVAILABLE:
            return self._create_failed_content("PDF libraries not available")
        
        start_time = time.time()
        
        # Try extraction methods in order of preference
        for method in self.available_methods:
            try:
                if method == "pdfplumber":
                    content = self._extract_with_pdfplumber(file_path)
                elif method == "pypdf2":
                    content = self._extract_with_pypdf2(file_path)
                elif method == "pdfminer":
                    content = self._extract_with_pdfminer(file_path)
                else:
                    continue
                
                if content and content.full_text.strip():
                    content.extraction_time = time.time() - start_time
                    content.metadata.text_extraction_method = method
                    return content
                    
            except Exception as e:
                logger.warning(f"PDF extraction method {method} failed: {e}")
                continue
        
        # If all methods failed
        return self._create_failed_content("All PDF extraction methods failed")
    
    def _extract_with_pdfplumber(self, file_path: Path) -> ExtractedContent:
        """Extract using pdfplumber (best for tables and layout)"""
        import pdfplumber
        
        content = ExtractedContent()
        full_text_parts = []
        
        with pdfplumber.open(file_path) as pdf:
            metadata = self._extract_pdf_metadata(file_path)
            metadata.page_count = len(pdf.pages)
            
            for page_num, page in enumerate(pdf.pages):
                # Extract text
                page_text = page.extract_text()
                if page_text:
                    full_text_parts.append(page_text)
                    
                    # Split into paragraphs
                    paragraphs = [p.strip() for p in page_text.split('\n\n') if p.strip()]
                    content.paragraphs.extend(paragraphs)
                
                # Extract tables
                tables = page.extract_tables()
                if tables:
                    metadata.has_tables = True
                    for table_idx, table in enumerate(tables):
                        if table:
                            content.tables.append({
                                "page": page_num + 1,
                                "table_index": table_idx,
                                "data": table,
                                "row_count": len(table),
                                "col_count": len(table[0]) if table else 0
                            })
                
                # Check for images
                if page.images:
                    metadata.has_images = True
                    for img_idx, img in enumerate(page.images):
                        content.images.append({
                            "page": page_num + 1,
                            "image_index": img_idx,
                            "bbox": img.get('bbox'),
                            "width": img.get('width'),
                            "height": img.get('height')
                        })
        
        content.full_text = '\n\n'.join(full_text_parts)
        content.metadata = metadata
        
        # Assess quality
        if content.full_text.strip():
            content.metadata.extraction_quality = ExtractionQuality.EXCELLENT
            content.metadata.extraction_confidence = 0.9
        
        return content
    
    def _extract_with_pypdf2(self, file_path: Path) -> ExtractedContent:
        """Extract using PyPDF2 (good for basic text)"""
        import PyPDF2
        
        content = ExtractedContent()
        full_text_parts = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            metadata = self._extract_pdf_metadata(file_path)
            metadata.page_count = len(pdf_reader.pages)
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        full_text_parts.append(page_text)
                        
                        # Split into paragraphs
                        paragraphs = [p.strip() for p in page_text.split('\n\n') if p.strip()]
                        content.paragraphs.extend(paragraphs)
                        
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num}: {e}")
        
        content.full_text = '\n\n'.join(full_text_parts)
        content.metadata = metadata
        
        # Assess quality
        if content.full_text.strip():
            content.metadata.extraction_quality = ExtractionQuality.GOOD
            content.metadata.extraction_confidence = 0.7
        
        return content
    
    def _extract_with_pdfminer(self, file_path: Path) -> ExtractedContent:
        """Extract using pdfminer (good for complex layouts)"""
        from pdfminer.high_level import extract_text
        
        content = ExtractedContent()
        
        try:
            full_text = extract_text(str(file_path))
            content.full_text = full_text
            
            # Split into paragraphs
            paragraphs = [p.strip() for p in full_text.split('\n\n') if p.strip()]
            content.paragraphs = paragraphs
            
            metadata = self._extract_pdf_metadata(file_path)
            content.metadata = metadata
            
            # Assess quality
            if content.full_text.strip():
                content.metadata.extraction_quality = ExtractionQuality.GOOD
                content.metadata.extraction_confidence = 0.8
            
        except Exception as e:
            logger.error(f"PDFMiner extraction failed: {e}")
            return self._create_failed_content(f"PDFMiner extraction failed: {e}")
        
        return content
    
    def _extract_pdf_metadata(self, file_path: Path) -> DocumentMetadata:
        """Extract PDF metadata"""
        metadata = DocumentMetadata(
            filename=file_path.name,
            file_path=str(file_path),
            file_size_bytes=file_path.stat().st_size,
            detected_format=DocumentFormat.PDF,
            mime_type="application/pdf"
        )
        
        # Calculate file hash
        metadata.file_hash = self._calculate_file_hash(file_path)
        
        # Try to extract PDF-specific metadata
        try:
            import PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                if pdf_reader.metadata:
                    pdf_meta = pdf_reader.metadata
                    metadata.title = pdf_meta.get('/Title')
                    metadata.author = pdf_meta.get('/Author')
                    metadata.subject = pdf_meta.get('/Subject')
                    metadata.creator = pdf_meta.get('/Creator')
                    metadata.producer = pdf_meta.get('/Producer')
                    
                    # Handle creation and modification dates
                    if '/CreationDate' in pdf_meta:
                        try:
                            creation_date_str = pdf_meta['/CreationDate']
                            # PDF dates are in format: D:YYYYMMDDHHmmSSOHH'mm'
                            if creation_date_str.startswith('D:'):
                                date_part = creation_date_str[2:16]  # YYYYMMDDHHmmSS
                                metadata.created_date = datetime.strptime(date_part, '%Y%m%d%H%M%S')
                        except Exception:
                            pass
        
        except Exception as e:
            logger.warning(f"Failed to extract PDF metadata: {e}")
        
        return metadata
    
    def _calculate_file_hash(self, file_path: Path) -> str:
        """Calculate SHA-256 hash of file"""
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()
    
    def _create_failed_content(self, error_message: str) -> ExtractedContent:
        """Create failed extraction content"""
        content = ExtractedContent()
        content.metadata.extraction_quality = ExtractionQuality.FAILED
        content.metadata.extraction_confidence = 0.0
        content.processing_notes.append(error_message)
        return content


class DocxExtractor(DocumentExtractor):
    """Microsoft Word DOCX extractor"""
    
    def can_extract(self, file_path: Path, mime_type: str) -> bool:
        """Check if this is a DOCX file"""
        return (DOCX_AVAILABLE and 
                (file_path.suffix.lower() == '.docx' or 
                 mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document']))
    
    def get_supported_formats(self) -> List[DocumentFormat]:
        """Get supported formats"""
        return [DocumentFormat.DOCX] if DOCX_AVAILABLE else []
    
    def extract_content(self, file_path: Path) -> ExtractedContent:
        """Extract content from DOCX file"""
        if not DOCX_AVAILABLE:
            return self._create_failed_content("python-docx library not available")
        
        start_time = time.time()
        
        try:
            doc = DocxDocument(str(file_path))
            content = ExtractedContent()
            
            # Extract metadata
            metadata = self._extract_docx_metadata(file_path, doc)
            
            # Extract text content
            full_text_parts = []
            
            # Process document elements in order
            for element in doc.element.body:
                if isinstance(element, CT_P):
                    # Paragraph
                    paragraph = Paragraph(element, doc)
                    para_text = paragraph.text.strip()
                    if para_text:
                        full_text_parts.append(para_text)
                        content.paragraphs.append(para_text)
                        
                        # Check if it's a heading
                        if paragraph.style.name.startswith('Heading'):
                            level = int(paragraph.style.name.split()[-1]) if paragraph.style.name.split()[-1].isdigit() else 1
                            content.headings.append({
                                "text": para_text,
                                "level": level,
                                "style": paragraph.style.name
                            })
                            metadata.has_headers = True
                
                elif isinstance(element, CT_Tbl):
                    # Table
                    table = Table(element, doc)
                    metadata.has_tables = True
                    
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        table_data.append(row_data)
                    
                    content.tables.append({
                        "data": table_data,
                        "row_count": len(table_data),
                        "col_count": len(table_data[0]) if table_data else 0
                    })
                    
                    # Add table text to full text
                    table_text = '\n'.join(['\t'.join(row) for row in table_data])
                    full_text_parts.append(table_text)
            
            # Check for images and links
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    metadata.has_images = True
                    content.images.append({
                        "relationship_id": rel.rId,
                        "target": rel.target_ref
                    })
            
            # Extract hyperlinks
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if run.element.xpath('.//w:hyperlink'):
                        metadata.has_links = True
                        # Extract hyperlink details would require more complex parsing
            
            content.full_text = '\n\n'.join(full_text_parts)
            content.metadata = metadata
            content.extraction_time = time.time() - start_time
            
            # Assess quality
            if content.full_text.strip():
                content.metadata.extraction_quality = ExtractionQuality.EXCELLENT
                content.metadata.extraction_confidence = 0.95
                content.metadata.text_extraction_method = "python-docx"
            
            return content
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return self._create_failed_content(f"DOCX extraction failed: {e}")
    
    def _extract_docx_metadata(self, file_path: Path, doc: DocxDocument) -> DocumentMetadata:
        """Extract DOCX metadata"""
        metadata = DocumentMetadata(
            filename=file_path.name,
            file_path=str(file_path),
            file_size_bytes=file_path.stat().st_size,
            detected_format=DocumentFormat.DOCX,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
        # Calculate file hash
        metadata.file_hash = self._calculate_file_hash(file_path)
        
        # Extract document properties
        try:
            core_props = doc.core_properties
            metadata.title = core_props.title
            metadata.author = core_props.author
            metadata.subject = core_props.subject
            metadata.creator = core_props.author
            metadata.created_date = core_props.created
            metadata.modified_date = core_props.modified
            
            if core_props.keywords:
                metadata.keywords = [k.strip() for k in core_props.keywords.split(',')]
        
        except Exception as e:
            logger.warning(f"Failed to extract DOCX properties: {e}")
        
        return metadata
    
    def _calculate_file_hash(self, file_path: Path) -> str:
        """Calculate SHA-256 hash of file"""
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()
    
    def _create_failed_content(self, error_message: str) -> ExtractedContent:
        """Create failed extraction content"""
        content = ExtractedContent()
        content.metadata.extraction_quality = ExtractionQuality.FAILED
        content.metadata.extraction_confidence = 0.0
        content.processing_notes.append(error_message)
        return content


class HTMLExtractor(DocumentExtractor):
    """HTML document extractor"""
    
    def can_extract(self, file_path: Path, mime_type: str) -> bool:
        """Check if this is an HTML file"""
        return (HTML_AVAILABLE and 
                (file_path.suffix.lower() in ['.html', '.htm'] or 
                 mime_type in ['text/html']))
    
    def get_supported_formats(self) -> List[DocumentFormat]:
        """Get supported formats"""
        return [DocumentFormat.HTML] if HTML_AVAILABLE else []
    
    def extract_content(self, file_path: Path) -> ExtractedContent:
        """Extract content from HTML file"""
        if not HTML_AVAILABLE:
            return self._create_failed_content("BeautifulSoup library not available")
        
        start_time = time.time()
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            content = ExtractedContent()
            
            # Extract metadata
            metadata = self._extract_html_metadata(file_path, soup)
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extract text content
            full_text = soup.get_text()
            
            # Clean up text
            lines = (line.strip() for line in full_text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            full_text = ' '.join(chunk for chunk in chunks if chunk)
            
            content.full_text = full_text
            
            # Extract paragraphs
            paragraphs = [p.get_text().strip() for p in soup.find_all('p') if p.get_text().strip()]
            content.paragraphs = paragraphs
            
            # Extract headings
            for i in range(1, 7):
                headings = soup.find_all(f'h{i}')
                for heading in headings:
                    heading_text = heading.get_text().strip()
                    if heading_text:
                        content.headings.append({
                            "text": heading_text,
                            "level": i,
                            "tag": f"h{i}"
                        })
                        metadata.has_headers = True
            
            # Extract tables
            tables = soup.find_all('table')
            if tables:
                metadata.has_tables = True
                for table_idx, table in enumerate(tables):
                    rows = table.find_all('tr')
                    table_data = []
                    for row in rows:
                        cells = row.find_all(['td', 'th'])
                        row_data = [cell.get_text().strip() for cell in cells]
                        if row_data:
                            table_data.append(row_data)
                    
                    if table_data:
                        content.tables.append({
                            "table_index": table_idx,
                            "data": table_data,
                            "row_count": len(table_data),
                            "col_count": len(table_data[0]) if table_data else 0
                        })
            
            # Extract links
            links = soup.find_all('a', href=True)
            if links:
                metadata.has_links = True
                for link in links:
                    link_text = link.get_text().strip()
                    href = link['href']
                    if link_text and href:
                        content.links.append({
                            "text": link_text,
                            "url": href,
                            "is_external": href.startswith(('http://', 'https://'))
                        })
            
            # Extract images
            images = soup.find_all('img', src=True)
            if images:
                metadata.has_images = True
                for img_idx, img in enumerate(images):
                    content.images.append({
                        "image_index": img_idx,
                        "src": img['src'],
                        "alt": img.get('alt', ''),
                        "title": img.get('title', '')
                    })
            
            content.metadata = metadata
            content.extraction_time = time.time() - start_time
            
            # Assess quality
            if content.full_text.strip():
                content.metadata.extraction_quality = ExtractionQuality.EXCELLENT
                content.metadata.extraction_confidence = 0.9
                content.metadata.text_extraction_method = "beautifulsoup"
            
            return content
            
        except Exception as e:
            logger.error(f"HTML extraction failed: {e}")
            return self._create_failed_content(f"HTML extraction failed: {e}")
    
    def _extract_html_metadata(self, file_path: Path, soup: BeautifulSoup) -> DocumentMetadata:
        """Extract HTML metadata"""
        metadata = DocumentMetadata(
            filename=file_path.name,
            file_path=str(file_path),
            file_size_bytes=file_path.stat().st_size,
            detected_format=DocumentFormat.HTML,
            mime_type="text/html"
        )
        
        # Calculate file hash
        metadata.file_hash = self._calculate_file_hash(file_path)
        
        # Extract HTML meta tags
        try:
            title_tag = soup.find('title')
            if title_tag:
                metadata.title = title_tag.get_text().strip()
            
            # Extract meta tags
            meta_tags = soup.find_all('meta')
            for meta in meta_tags:
                name = meta.get('name', '').lower()
                content = meta.get('content', '')
                
                if name == 'author':
                    metadata.author = content
                elif name == 'description':
                    metadata.subject = content
                elif name == 'keywords':
                    metadata.keywords = [k.strip() for k in content.split(',')]
                elif name == 'generator':
                    metadata.creator = content
        
        except Exception as e:
            logger.warning(f"Failed to extract HTML metadata: {e}")
        
        return metadata
    
    def _calculate_file_hash(self, file_path: Path) -> str:
        """Calculate SHA-256 hash of file"""
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()
    
    def _create_failed_content(self, error_message: str) -> ExtractedContent:
        """Create failed extraction content"""
        content = ExtractedContent()
        content.metadata.extraction_quality = ExtractionQuality.FAILED
        content.metadata.extraction_confidence = 0.0
        content.processing_notes.append(error_message)
        return content


class TextExtractor(DocumentExtractor):
    """Plain text and fallback extractor"""
    
    def can_extract(self, file_path: Path, mime_type: str) -> bool:
        """Check if this is a text file or unknown format"""
        text_extensions = ['.txt', '.text', '.log', '.md', '.markdown', '.csv', '.json', '.xml']
        text_mimes = ['text/plain', 'text/csv', 'application/json', 'text/xml', 'application/xml']
        
        return (file_path.suffix.lower() in text_extensions or 
                mime_type in text_mimes or
                mime_type.startswith('text/'))
    
    def get_supported_formats(self) -> List[DocumentFormat]:
        """Get supported formats"""
        return [DocumentFormat.TXT, DocumentFormat.CSV, DocumentFormat.JSON, 
                DocumentFormat.XML, DocumentFormat.MARKDOWN]
    
    def extract_content(self, file_path: Path) -> ExtractedContent:
        """Extract content from text file"""
        start_time = time.time()
        
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            text_content = None
            used_encoding = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text_content = f.read()
                    used_encoding = encoding
                    break
                except (UnicodeDecodeError, UnicodeError):
                    continue
            
            if text_content is None:
                return self._create_failed_content("Could not decode file with any supported encoding")
            
            content = ExtractedContent()
            
            # Extract metadata
            metadata = self._extract_text_metadata(file_path, text_content, used_encoding)
            
            content.full_text = text_content
            
            # Split into paragraphs
            paragraphs = [p.strip() for p in text_content.split('\n\n') if p.strip()]
            content.paragraphs = paragraphs
            
            # Detect format-specific content
            file_ext = file_path.suffix.lower()
            
            if file_ext == '.json':
                content = self._extract_json_content(content, text_content)
                metadata.detected_format = DocumentFormat.JSON
            elif file_ext == '.csv':
                content = self._extract_csv_content(content, text_content)
                metadata.detected_format = DocumentFormat.CSV  
            elif file_ext in ['.md', '.markdown']:
                content = self._extract_markdown_content(content, text_content)
                metadata.detected_format = DocumentFormat.MARKDOWN
            elif file_ext == '.xml':
                content = self._extract_xml_content(content, text_content)
                metadata.detected_format = DocumentFormat.XML
            
            content.metadata = metadata
            content.extraction_time = time.time() - start_time
            
            # Assess quality
            if content.full_text.strip():
                content.metadata.extraction_quality = ExtractionQuality.EXCELLENT
                content.metadata.extraction_confidence = 1.0
                content.metadata.text_extraction_method = f"text_file_{used_encoding}"
            
            return content
            
        except Exception as e:
            logger.error(f"Text extraction failed: {e}")
            return self._create_failed_content(f"Text extraction failed: {e}")
    
    def _extract_text_metadata(self, file_path: Path, content: str, encoding: str) -> DocumentMetadata:
        """Extract text file metadata"""
        metadata = DocumentMetadata(
            filename=file_path.name,
            file_path=str(file_path),
            file_size_bytes=file_path.stat().st_size,
            detected_format=DocumentFormat.TXT,
            mime_type="text/plain",
            encoding=encoding
        )
        
        # Calculate file hash
        metadata.file_hash = self._calculate_file_hash(file_path)
        
        # Basic text statistics
        metadata.character_count = len(content)
        metadata.word_count = len(content.split())
        metadata.paragraph_count = len([p for p in content.split('\n\n') if p.strip()])
        
        # File timestamps
        try:
            stat = file_path.stat()
            metadata.created_date = datetime.fromtimestamp(stat.st_ctime)
            metadata.modified_date = datetime.fromtimestamp(stat.st_mtime)
            metadata.accessed_date = datetime.fromtimestamp(stat.st_atime)
        except Exception:
            pass
        
        return metadata
    
    def _extract_json_content(self, content: ExtractedContent, text: str) -> ExtractedContent:
        """Extract structured content from JSON"""
        try:
            json_data = json.loads(text)
            content.structured_text = {"json_data": json_data}
            content.metadata.has_headers = True  # JSON has structure
        except json.JSONDecodeError as e:
            content.processing_notes.append(f"JSON parsing failed: {e}")
        
        return content
    
    def _extract_csv_content(self, content: ExtractedContent, text: str) -> ExtractedContent:
        """Extract structured content from CSV"""
        try:
            lines = text.strip().split('\n')
            if lines:
                # Assume first line is header
                if len(lines) > 1:
                    content.metadata.has_headers = True
                    content.metadata.has_tables = True
                    
                    # Simple CSV parsing (could use csv module for better parsing)
                    table_data = []
                    for line in lines:
                        row = [cell.strip() for cell in line.split(',')]
                        table_data.append(row)
                    
                    content.tables.append({
                        "data": table_data,
                        "row_count": len(table_data),
                        "col_count": len(table_data[0]) if table_data else 0,
                        "has_header": True
                    })
        except Exception as e:
            content.processing_notes.append(f"CSV parsing failed: {e}")
        
        return content
    
    def _extract_markdown_content(self, content: ExtractedContent, text: str) -> ExtractedContent:
        """Extract structured content from Markdown"""
        try:
            # Extract headers
            lines = text.split('\n')
            for line in lines:
                line = line.strip()
                if line.startswith('#'):
                    level = len(line) - len(line.lstrip('#'))
                    header_text = line.lstrip('#').strip()
                    if header_text:
                        content.headings.append({
                            "text": header_text,
                            "level": level,
                            "markdown_level": level
                        })
                        content.metadata.has_headers = True
            
            # Extract links
            link_pattern = r'\[([^\]]+)\]\(([^\)]+)\)'
            links = re.findall(link_pattern, text)
            if links:
                content.metadata.has_links = True
                for link_text, url in links:
                    content.links.append({
                        "text": link_text,
                        "url": url,
                        "is_external": url.startswith(('http://', 'https://'))
                    })
            
            # Extract images
            image_pattern = r'!\[([^\]]*)\]\(([^\)]+)\)'
            images = re.findall(image_pattern, text)
            if images:
                content.metadata.has_images = True
                for img_idx, (alt_text, src) in enumerate(images):
                    content.images.append({
                        "image_index": img_idx,
                        "src": src,
                        "alt": alt_text
                    })
            
        except Exception as e:
            content.processing_notes.append(f"Markdown parsing failed: {e}")
        
        return content
    
    def _extract_xml_content(self, content: ExtractedContent, text: str) -> ExtractedContent:
        """Extract structured content from XML"""
        try:
            if HTML_AVAILABLE:
                soup = BeautifulSoup(text, 'xml')
                
                # Get text content
                text_content = soup.get_text()
                content.full_text = text_content
                
                # XML has inherent structure
                content.metadata.has_headers = True
                
        except Exception as e:
            content.processing_notes.append(f"XML parsing failed: {e}")
        
        return content
    
    def _calculate_file_hash(self, file_path: Path) -> str:
        """Calculate SHA-256 hash of file"""
        sha256_hash = hashlib.sha256()
        with open(file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        return sha256_hash.hexdigest()
    
    def _create_failed_content(self, error_message: str) -> ExtractedContent:
        """Create failed extraction content"""
        content = ExtractedContent()
        content.metadata.extraction_quality = ExtractionQuality.FAILED
        content.metadata.extraction_confidence = 0.0
        content.processing_notes.append(error_message)
        return content


class DocumentIngestionManager:
    """Main document ingestion coordinator"""
    
    def __init__(self):
        # Initialize extractors
        self.extractors = [
            PDFExtractor(),
            DocxExtractor(), 
            HTMLExtractor(),
            TextExtractor()  # Keep as fallback
        ]
        
        # Filter available extractors
        self.available_extractors = [
            extractor for extractor in self.extractors
            if extractor.get_supported_formats()
        ]
        
        # Statistics
        self.ingestion_stats = {
            'total_documents': 0,
            'successful_extractions': 0,
            'failed_extractions': 0,
            'total_extraction_time': 0.0,
            'formats_processed': defaultdict(int),
            'start_time': datetime.now()
        }
        
        logger.info(f"DocumentIngestionManager initialized with {len(self.available_extractors)} extractors")
    
    def detect_format(self, file_path: Path) -> Tuple[DocumentFormat, str]:
        """Detect document format and MIME type"""
        try:
            # Get MIME type
            mime_type, _ = mimetypes.guess_type(str(file_path))
            if not mime_type:
                mime_type = "application/octet-stream"
            
            # Map file extension to format
            ext = file_path.suffix.lower()
            format_mapping = {
                '.pdf': DocumentFormat.PDF,
                '.docx': DocumentFormat.DOCX,
                '.doc': DocumentFormat.DOC,
                '.pptx': DocumentFormat.PPTX,
                '.ppt': DocumentFormat.PPT,
                '.xlsx': DocumentFormat.XLSX,
                '.xls': DocumentFormat.XLS,
                '.csv': DocumentFormat.CSV,
                '.html': DocumentFormat.HTML,
                '.htm': DocumentFormat.HTML,
                '.xml': DocumentFormat.XML,
                '.md': DocumentFormat.MARKDOWN,
                '.markdown': DocumentFormat.MARKDOWN,
                '.txt': DocumentFormat.TXT,
                '.json': DocumentFormat.JSON,
                '.rtf': DocumentFormat.RTF,
                '.eml': DocumentFormat.EMAIL,
                '.epub': DocumentFormat.EPUB,
                '.odt': DocumentFormat.ODT
            }
            
            detected_format = format_mapping.get(ext, DocumentFormat.UNKNOWN)
            
            return detected_format, mime_type
            
        except Exception as e:
            logger.warning(f"Format detection failed for {file_path}: {e}")
            return DocumentFormat.UNKNOWN, "application/octet-stream"
    
    def find_suitable_extractor(self, file_path: Path, mime_type: str) -> Optional[DocumentExtractor]:
        """Find suitable extractor for file"""
        for extractor in self.available_extractors:
            if extractor.can_extract(file_path, mime_type):
                return extractor
        return None
    
    def ingest_document(self, file_path: Union[str, Path]) -> ExtractedContent:
        """Ingest single document"""
        start_time = time.time()
        file_path = Path(file_path)
        
        try:
            if not file_path.exists():
                return self._create_failed_content(f"File not found: {file_path}")
            
            # Detect format
            detected_format, mime_type = self.detect_format(file_path)
            
            # Find suitable extractor
            extractor = self.find_suitable_extractor(file_path, mime_type)
            
            if not extractor:
                return self._create_failed_content(f"No suitable extractor found for {file_path}")
            
            # Extract content
            logger.info(f"Extracting content from {file_path.name} using {extractor.__class__.__name__}")
            
            content = extractor.extract_content(file_path)
            
            # Update statistics
            self.ingestion_stats['total_documents'] += 1
            self.ingestion_stats['formats_processed'][detected_format.value] += 1
            
            if content.metadata.extraction_quality != ExtractionQuality.FAILED:
                self.ingestion_stats['successful_extractions'] += 1
            else:
                self.ingestion_stats['failed_extractions'] += 1
            
            total_time = time.time() - start_time
            self.ingestion_stats['total_extraction_time'] += total_time
            
            # Enhance metadata
            if not content.metadata.filename:
                content.metadata.filename = file_path.name
            if not content.metadata.file_path:
                content.metadata.file_path = str(file_path)
            if not content.metadata.detected_format or content.metadata.detected_format == DocumentFormat.UNKNOWN:
                content.metadata.detected_format = detected_format
            if not content.metadata.mime_type:
                content.metadata.mime_type = mime_type
            
            # Add word and character counts if not set
            if content.full_text:
                content.metadata.word_count = len(content.full_text.split())
                content.metadata.character_count = len(content.full_text)
                content.metadata.paragraph_count = len(content.paragraphs)
            
            logger.info(f"Successfully extracted {len(content.full_text)} characters from {file_path.name}")
            
            return content
            
        except Exception as e:
            logger.error(f"Document ingestion failed for {file_path}: {e}")
            self.ingestion_stats['total_documents'] += 1
            self.ingestion_stats['failed_extractions'] += 1
            return self._create_failed_content(f"Ingestion failed: {e}")
    
    def ingest_batch(self, file_paths: List[Union[str, Path]], 
                    max_workers: int = 4) -> List[ExtractedContent]:
        """Ingest multiple documents in parallel"""
        from concurrent.futures import ThreadPoolExecutor, as_completed
        
        results = []
        
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            # Submit all ingestion tasks
            future_to_path = {
                executor.submit(self.ingest_document, path): path 
                for path in file_paths
            }
            
            # Collect results as they complete
            for future in as_completed(future_to_path):
                path = future_to_path[future]
                try:
                    content = future.result()
                    results.append(content)
                except Exception as e:
                    logger.error(f"Batch ingestion failed for {path}: {e}")
                    results.append(self._create_failed_content(f"Batch processing failed: {e}"))
        
        logger.info(f"Batch ingestion completed: {len(results)} documents processed")
        
        return results
    
    def ingest_directory(self, directory_path: Union[str, Path], 
                        recursive: bool = True,
                        file_patterns: Optional[List[str]] = None,
                        max_workers: int = 4) -> List[ExtractedContent]:
        """Ingest all documents from directory"""
        directory_path = Path(directory_path)
        
        if not directory_path.exists() or not directory_path.is_dir():
            logger.error(f"Directory not found: {directory_path}")
            return []
        
        # Find all files
        if recursive:
            files = list(directory_path.rglob('*'))
        else:
            files = list(directory_path.glob('*'))
        
        # Filter to files only
        files = [f for f in files if f.is_file()]
        
        # Apply file patterns if specified
        if file_patterns:
            filtered_files = []
            for pattern in file_patterns:
                if recursive:
                    pattern_files = list(directory_path.rglob(pattern))
                else:
                    pattern_files = list(directory_path.glob(pattern))
                filtered_files.extend([f for f in pattern_files if f.is_file()])
            files = list(set(filtered_files))  # Remove duplicates
        
        logger.info(f"Found {len(files)} files in {directory_path}")
        
        # Ingest all files
        return self.ingest_batch(files, max_workers=max_workers)
    
    def get_ingestion_stats(self) -> Dict[str, Any]:
        """Get ingestion statistics"""
        total_time = (datetime.now() - self.ingestion_stats['start_time']).total_seconds()
        
        stats = dict(self.ingestion_stats)
        stats['formats_processed'] = dict(stats['formats_processed'])
        stats['total_time_seconds'] = total_time
        
        if stats['total_documents'] > 0:
            stats['success_rate'] = stats['successful_extractions'] / stats['total_documents']
            stats['average_extraction_time'] = stats['total_extraction_time'] / stats['total_documents']
        else:
            stats['success_rate'] = 0.0
            stats['average_extraction_time'] = 0.0
        
        stats['available_extractors'] = [
            extractor.get_extractor_info() for extractor in self.available_extractors
        ]
        
        return stats
    
    def get_supported_formats(self) -> List[str]:
        """Get list of all supported formats"""
        supported = set()
        for extractor in self.available_extractors:
            for fmt in extractor.get_supported_formats():
                supported.add(fmt.value)
        return sorted(list(supported))
    
    def export_extraction_results(self, contents: List[ExtractedContent], 
                                export_path: str) -> bool:
        """Export extraction results to JSON file"""
        try:
            export_data = {
                "export_timestamp": datetime.now().isoformat(),
                "total_documents": len(contents),
                "ingestion_stats": self.get_ingestion_stats(),
                "documents": []
            }
            
            for content in contents:
                doc_data = {
                    "metadata": asdict(content.metadata),
                    "content_summary": {
                        "text_length": len(content.full_text),
                        "paragraph_count": len(content.paragraphs),
                        "heading_count": len(content.headings),
                        "table_count": len(content.tables),
                        "image_count": len(content.images),
                        "link_count": len(content.links)
                    },
                    "extraction_time": content.extraction_time,
                    "processing_notes": content.processing_notes,
                    # Include first 1000 characters of text for preview
                    "text_preview": content.full_text[:1000] if content.full_text else ""
                }
                
                export_data["documents"].append(doc_data)
            
            # Write to file
            export_path = Path(export_path)
            export_path.parent.mkdir(parents=True, exist_ok=True)
            
            with open(export_path, 'w', encoding='utf-8') as f:
                json.dump(export_data, f, indent=2, default=str, ensure_ascii=False)
            
            logger.info(f"Exported extraction results to {export_path}")
            return True
            
        except Exception as e:
            logger.error(f"Failed to export extraction results: {e}")
            return False
    
    def _create_failed_content(self, error_message: str) -> ExtractedContent:
        """Create failed extraction content"""
        content = ExtractedContent()
        content.metadata.extraction_quality = ExtractionQuality.FAILED
        content.metadata.extraction_confidence = 0.0
        content.processing_notes.append(error_message)
        return content


# Factory functions for common use cases
def create_document_ingestion_manager() -> DocumentIngestionManager:
    """Create document ingestion manager with all available extractors"""
    return DocumentIngestionManager()


# Context manager for temporary file processing
@contextmanager
def temporary_file_processing(file_data: bytes, file_extension: str):
    """Context manager for processing files from memory"""
    with tempfile.NamedTemporaryFile(suffix=file_extension, delete=False) as temp_file:
        temp_file.write(file_data)
        temp_path = Path(temp_file.name)
    
    try:
        yield temp_path
    finally:
        try:
            temp_path.unlink()
        except Exception as e:
            logger.warning(f"Failed to cleanup temporary file {temp_path}: {e}")


# Example usage and testing
if __name__ == "__main__":
    def test_document_ingestion():
        # Create ingestion manager
        manager = create_document_ingestion_manager()
        
        print("Document Ingestion System Test")
        print("=" * 50)
        
        # Show supported formats
        supported_formats = manager.get_supported_formats()
        print(f"Supported formats: {', '.join(supported_formats)}")
        print()
        
        # Show available extractors
        stats = manager.get_ingestion_stats()
        print("Available extractors:")
        for extractor_info in stats['available_extractors']:
            print(f"  - {extractor_info['extractor_name']}: {', '.join(extractor_info['supported_formats'])}")
        print()
        
        # Test with some sample files (if they exist)
        test_files = [
            "test_data/sample.pdf",
            "test_data/sample.docx", 
            "test_data/sample.html",
            "test_data/sample.txt"
        ]
        
        existing_files = [f for f in test_files if Path(f).exists()]
        
        if existing_files:
            print(f"Testing with {len(existing_files)} files...")
            
            # Ingest documents
            results = manager.ingest_batch(existing_files)
            
            print("\nIngestion Results:")
            for i, content in enumerate(results):
                metadata = content.metadata
                print(f"\n{i+1}. {metadata.filename}")
                print(f"   Format: {metadata.detected_format.value}")
                print(f"   Quality: {metadata.extraction_quality.value}")
                print(f"   Text length: {len(content.full_text)} characters")
                print(f"   Paragraphs: {len(content.paragraphs)}")
                print(f"   Tables: {len(content.tables)}")
                print(f"   Images: {len(content.images)}")
                
                if content.processing_notes:
                    print(f"   Notes: {'; '.join(content.processing_notes)}")
            
            # Show final statistics
            final_stats = manager.get_ingestion_stats()
            print(f"\nFinal Statistics:")
            print(f"  Total documents: {final_stats['total_documents']}")
            print(f"  Successful: {final_stats['successful_extractions']}")
            print(f"  Failed: {final_stats['failed_extractions']}")
            print(f"  Success rate: {final_stats['success_rate']:.1%}")
            print(f"  Average time: {final_stats['average_extraction_time']:.2f}s")
            
            # Export results
            if manager.export_extraction_results(results, "test_extraction_results.json"):
                print("\nResults exported to test_extraction_results.json")
        
        else:
            print("No test files found. Create some sample files to test extraction.")
            print(f"Expected files: {', '.join(test_files)}")
    
    test_document_ingestion()