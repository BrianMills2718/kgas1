"""
T02: Word Document Loader - Unified Interface Implementation

Loads and extracts text from Word documents (.docx) using python-docx.
"""

from typing import Dict, Any, Optional, List
import os
from pathlib import Path
import uuid
from datetime import datetime
import logging

from src.tools.base_tool import BaseTool, ToolRequest, ToolResult, ToolContract, ToolStatus
from src.core.service_manager import ServiceManager

logger = logging.getLogger(__name__)


class T02WordLoaderUnified(BaseTool):
    """T02: Word Document Loader with unified interface"""
    
    def __init__(self, service_manager: ServiceManager):
        """Initialize with service manager"""
        super().__init__(service_manager)
        self.tool_id = "T02_WORD_LOADER"
        self.identity_service = service_manager.identity_service
        self.provenance_service = service_manager.provenance_service
        self.quality_service = service_manager.quality_service
        self._temp_files = []
    
    def get_contract(self) -> ToolContract:
        """Return tool contract specification"""
        return ToolContract(
            tool_id=self.tool_id,
            name="Word Document Loader",
            description="Load and extract text from Word documents (.docx)",
            category="document_processing",
            input_schema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "Path to Word document to load"
                    },
                    "workflow_id": {
                        "type": "string",
                        "description": "Optional workflow ID for tracking"
                    }
                },
                "required": ["file_path"]
            },
            output_schema={
                "type": "object",
                "properties": {
                    "document": {
                        "type": "object",
                        "properties": {
                            "document_id": {"type": "string"},
                            "document_ref": {"type": "string"},
                            "file_path": {"type": "string"},
                            "file_name": {"type": "string"},
                            "file_size": {"type": "integer"},
                            "paragraph_count": {"type": "integer"},
                            "table_count": {"type": "integer"},
                            "text": {"type": "string"},
                            "text_length": {"type": "integer"},
                            "confidence": {"type": "number"},
                            "quality_tier": {"type": "string"},
                            "created_at": {"type": "string"}
                        },
                        "required": ["document_id", "text", "confidence", "paragraph_count"]
                    }
                },
                "required": ["document"]
            },
            dependencies=["identity_service", "provenance_service", "quality_service"],
            performance_requirements={
                "max_execution_time": 20.0,  # 20 seconds for large documents
                "max_memory_mb": 1024,       # 1GB for document processing
                "min_confidence": 0.8        # Minimum confidence threshold
            },
            error_conditions=[
                "FILE_NOT_FOUND",
                "INVALID_FILE_TYPE",
                "DOCX_CORRUPTED",
                "DOCX_PROTECTED",
                "EXTRACTION_FAILED",
                "MEMORY_LIMIT_EXCEEDED"
            ]
        )
    
    def execute(self, request: ToolRequest) -> ToolResult:
        """Execute Word document loading with unified interface"""
        self._start_execution()
        
        try:
            # Validate input
            if not self.validate_input(request.input_data):
                return self._create_error_result(
                    request,
                    "INVALID_INPUT",
                    "Input validation failed. Required: file_path"
                )
            
            # Extract parameters
            file_path = request.input_data.get("file_path")
            workflow_id = request.input_data.get("workflow_id")
            
            # Validate file path
            validation_result = self._validate_file_path(file_path)
            if not validation_result["valid"]:
                return self._create_error_result(
                    request,
                    validation_result["error_code"],
                    validation_result["error_message"]
                )
            
            file_path = Path(file_path)
            
            # Start provenance tracking
            operation_id = self.provenance_service.start_operation(
                tool_id=self.tool_id,
                operation_type="load_document",
                used={},
                parameters={
                    "file_path": str(file_path),
                    "workflow_id": workflow_id
                }
            )
            
            # Generate workflow ID if not provided
            if not workflow_id:
                workflow_id = f"wf_{uuid.uuid4().hex[:8]}"
            
            # Create document ID
            document_id = f"{workflow_id}_{file_path.stem}"
            document_ref = f"storage://document/{document_id}"
            
            # Extract text from Word document
            extraction_result = self._extract_text_from_docx(file_path, request.parameters)
            
            if extraction_result["status"] != "success":
                return self._create_error_result(
                    request,
                    extraction_result.get("error_code", "EXTRACTION_FAILED"),
                    extraction_result["error"]
                )
            
            # Calculate confidence
            confidence = self._calculate_confidence(
                text=extraction_result["text"],
                paragraph_count=extraction_result["paragraph_count"],
                table_count=extraction_result["table_count"],
                file_size=file_path.stat().st_size
            )
            
            # Create document data
            document_data = {
                "document_id": document_id,
                "document_ref": document_ref,
                "file_path": str(file_path),
                "file_name": file_path.name,
                "file_size": file_path.stat().st_size,
                "paragraph_count": extraction_result["paragraph_count"],
                "table_count": extraction_result["table_count"],
                "text": extraction_result["text"],
                "text_length": len(extraction_result["text"]),
                "confidence": confidence,
                "created_at": datetime.now().isoformat(),
                "tool_version": "1.0.0",
                "extraction_method": "python-docx"
            }
            
            # Assess quality
            quality_result = self.quality_service.assess_confidence(
                object_ref=document_ref,
                base_confidence=confidence,
                factors={
                    "text_length": min(1.0, len(extraction_result["text"]) / 10000),
                    "paragraph_count": min(1.0, extraction_result["paragraph_count"] / 50),
                    "table_count": min(1.0, extraction_result["table_count"] / 10) if extraction_result["table_count"] > 0 else 0.8,
                    "file_size": min(1.0, file_path.stat().st_size / (1024 * 1024))
                },
                metadata={
                    "extraction_method": document_data["extraction_method"],
                    "file_type": file_path.suffix.lower()
                }
            )
            
            if quality_result["status"] == "success":
                document_data["confidence"] = quality_result["confidence"]
                document_data["quality_tier"] = quality_result["quality_tier"]
            
            # Complete provenance
            self.provenance_service.complete_operation(
                operation_id=operation_id,
                outputs=[document_ref],
                success=True,
                metadata={
                    "paragraph_count": extraction_result["paragraph_count"],
                    "table_count": extraction_result["table_count"],
                    "text_length": len(extraction_result["text"]),
                    "confidence": document_data["confidence"]
                }
            )
            
            # Get execution metrics
            execution_time, memory_used = self._end_execution()
            
            # Create success result
            return ToolResult(
                tool_id=self.tool_id,
                status="success",
                data={
                    "document": document_data
                },
                metadata={
                    "operation_id": operation_id,
                    "workflow_id": workflow_id,
                    "extraction_method": document_data["extraction_method"]
                },
                execution_time=execution_time,
                memory_used=memory_used
            )
            
        except Exception as e:
            logger.error(f"Unexpected error in {self.tool_id}: {e}", exc_info=True)
            return self._create_error_result(
                request,
                "UNEXPECTED_ERROR",
                f"Unexpected error during Word document loading: {str(e)}"
            )
    
    def _validate_file_path(self, file_path: str) -> Dict[str, Any]:
        """Validate file path for security and existence"""
        if not file_path:
            return {
                "valid": False,
                "error_code": "INVALID_INPUT",
                "error_message": "File path cannot be empty"
            }
        
        try:
            path = Path(file_path)
            
            # Check if path exists
            if not path.exists():
                return {
                    "valid": False,
                    "error_code": "FILE_NOT_FOUND",
                    "error_message": f"File not found: {file_path}"
                }
            
            # Check if it's a file
            if not path.is_file():
                return {
                    "valid": False,
                    "error_code": "INVALID_INPUT",
                    "error_message": f"Path is not a file: {file_path}"
                }
            
            # Check extension
            allowed_extensions = ['.docx']
            if path.suffix.lower() not in allowed_extensions:
                return {
                    "valid": False,
                    "error_code": "INVALID_FILE_TYPE",
                    "error_message": f"Invalid file extension. Allowed: {allowed_extensions}"
                }
            
            # Basic security check - prevent path traversal
            if ".." in str(path) or str(path).startswith("/etc"):
                return {
                    "valid": False,
                    "error_code": "VALIDATION_FAILED",
                    "error_message": "Invalid file path"
                }
            
            return {"valid": True}
            
        except Exception as e:
            return {
                "valid": False,
                "error_code": "VALIDATION_FAILED",
                "error_message": f"Path validation failed: {str(e)}"
            }
    
    def _extract_text_from_docx(self, file_path: Path, parameters: Dict[str, Any]) -> Dict[str, Any]:
        """Extract text from Word document using python-docx"""
        try:
            import docx
            
            # Open the document
            doc = docx.Document(str(file_path))
            
            # Extract paragraphs
            paragraphs = []
            paragraph_count = 0
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:  # Skip empty paragraphs
                    paragraphs.append(text)
                    paragraph_count += 1
            
            # Extract tables if requested
            table_count = len(doc.tables)
            table_text = []
            
            if parameters.get("extract_tables", True):
                for table in doc.tables:
                    table_rows = []
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_text.append(cell_text)
                        if row_text:
                            table_rows.append(" | ".join(row_text))
                    if table_rows:
                        table_text.append("\n".join(table_rows))
            
            # Combine all text
            all_text = []
            all_text.extend(paragraphs)
            if table_text:
                all_text.append("\n\n[Tables]\n")
                all_text.extend(table_text)
            
            full_text = "\n\n".join(all_text)
            
            # Basic text cleaning
            cleaned_text = self._clean_extracted_text(full_text)
            
            return {
                "status": "success",
                "text": cleaned_text,
                "paragraph_count": paragraph_count,
                "table_count": table_count
            }
            
        except Exception as e:
            error_message = str(e).lower()
            
            # Determine specific error type
            if "package not found" in error_message:
                error_code = "DOCX_CORRUPTED"
            elif "password" in error_message or "protected" in error_message:
                error_code = "DOCX_PROTECTED"
            else:
                error_code = "EXTRACTION_FAILED"
            
            logger.error(f"Failed to extract text from Word document: {str(e)}")
            return {
                "status": "error",
                "error": f"Failed to extract text from Word document: {str(e)}",
                "error_code": error_code
            }
    
    def _clean_extracted_text(self, text: str) -> str:
        """Basic text cleaning for extracted text"""
        if not text:
            return ""
        
        import re
        
        # Replace multiple spaces with single space
        text = re.sub(r' +', ' ', text)
        
        # Replace multiple newlines with double newlines
        text = re.sub(r'\n\n+', '\n\n', text)
        
        # Remove leading/trailing whitespace from lines
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)
        
        # Remove empty lines at start and end
        text = text.strip()
        
        return text
    
    def _calculate_confidence(self, text: str, paragraph_count: int, table_count: int, file_size: int) -> float:
        """Calculate confidence score for extracted text"""
        base_confidence = 0.9  # High confidence for python-docx extraction
        
        # Factors that affect confidence
        factors = []
        
        # Text length factor
        if len(text) > 1000:
            factors.append(0.95)
        elif len(text) > 100:
            factors.append(0.85)
        else:
            factors.append(0.6)
        
        # Paragraph count factor
        if paragraph_count > 10:
            factors.append(0.95)
        elif paragraph_count > 3:
            factors.append(0.9)
        else:
            factors.append(0.8)
        
        # Table factor (bonus for structured content)
        if table_count > 0:
            factors.append(0.95)
        else:
            factors.append(0.9)
        
        # File size factor
        if file_size > 1024 * 1024:  # > 1MB
            factors.append(0.95)
        elif file_size > 100 * 1024:  # > 100KB
            factors.append(0.9)
        else:
            factors.append(0.8)
        
        # Calculate weighted average
        if factors:
            final_confidence = (base_confidence + sum(factors)) / (len(factors) + 1)
        else:
            final_confidence = base_confidence
        
        # Ensure confidence is in valid range
        return max(0.1, min(1.0, final_confidence))
    
    def health_check(self) -> ToolResult:
        """Check tool health and readiness"""
        try:
            # Check if python-docx is available
            try:
                import docx
                docx_available = True
            except ImportError:
                docx_available = False
            
            # Check service dependencies
            services_healthy = True
            if self.services:
                try:
                    # Basic check that services exist
                    _ = self.identity_service
                    _ = self.provenance_service
                    _ = self.quality_service
                except:
                    services_healthy = False
            
            healthy = docx_available and services_healthy
            
            return ToolResult(
                tool_id=self.tool_id,
                status="success" if healthy else "error",
                data={
                    "healthy": healthy,
                    "docx_available": docx_available,
                    "services_healthy": services_healthy,
                    "supported_formats": [".docx"],
                    "status": self.status.value
                },
                metadata={
                    "timestamp": datetime.now().isoformat()
                },
                execution_time=0.0,
                memory_used=0
            )
            
        except Exception as e:
            return ToolResult(
                tool_id=self.tool_id,
                status="error",
                data={"healthy": False},
                metadata={"error": str(e)},
                execution_time=0.0,
                memory_used=0,
                error_code="HEALTH_CHECK_FAILED",
                error_message=str(e)
            )
    
    def cleanup(self) -> bool:
        """Clean up any temporary files"""
        try:
            # Clean up temp files if any
            for temp_file in self._temp_files:
                try:
                    if os.path.exists(temp_file):
                        os.remove(temp_file)
                except:
                    pass
            
            self._temp_files = []
            self.status = ToolStatus.READY
            return True
            
        except Exception as e:
            logger.error(f"Cleanup failed: {e}")
            return False